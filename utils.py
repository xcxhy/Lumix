import os
import re
import json
import time
import pdfplumber
import pytesseract
from bs4 import BeautifulSoup
import cv2 as cv
import fitz
from PIL import Image
import shutil
import docx
import ocrmypdf
import math
import pexpect
import pickle
import subprocess
import shlex
from pypdf import PdfReader, PdfWriter
import random
import wordninja
from tqdm import tqdm
from io import BytesIO, StringIO
from requests import post
from datetime import datetime
from transformers import LlamaTokenizer
from sentence_transformers import SentenceTransformer, util
import subprocess
import shlex
from itertools import chain

# 读取json文件
def read_json(path):
    # new_data = []
    data = json.load(open(path, 'r', encoding='utf-8'))
    return data
# 保存json文件
def write_json(path, dataset):
    dict_json = json.dumps(dataset, indent=4, ensure_ascii=False)
    with open(path, "w", encoding='utf-8') as file:
        file.write(dict_json)
    return 
def read_file_or_dir(path):
    # start read file
    # judge file or dir
    if os.path.isfile(path):
        data = read_json(path)
        # get part data
        # data = data[:1000]
        cate = choose_file_category(path)
        return data, cate
    elif os.path.isdir(path):
        data = []
        files = os.listdir(path)
        for file in tqdm(files):
            if file.endswith(".json") or file.endswith(".jsonl"):
                file_path = os.path.join(path, file)
                value = read_json(file_path)
                data.extend(value)
        cate = choose_file_category(os.path.join(path, files[0]))
        return data, cate
    else:
        raise ValueError("path is not a file or dir")
    
# 增加json文件
def write_json_add(path, dataset):
    dict_json = json.dumps(dataset, indent=4, ensure_ascii=False)
    with open(path, "a", encoding='utf-8') as file:
        file.write(dict_json)
    return 
# 读取jsonl文件
def read_jsonl(path):
    all_data = []
    with open(path, 'r', encoding="utf-8") as f:
        for line in f:
            try:
                data = json.loads(line)
                all_data.append(data)
            except:
                pass
    return all_data
# 保存jsonl文件
def write_jsonl(path, dataset):
    with open(path, 'w', encoding="utf-8") as f:
        for data in dataset:
            f.write(json.dumps(data) + "\n")
    return 
# 添加新的jsonl数据
def write_jsonl_add(path, dataset):
    with open(path, 'a', encoding="utf-8") as f:
        for data in dataset:
            f.write(json.dumps(data) + "\n")
    return
# 获取时间戳
def get_timestamp() -> str:
    return datetime.now().isoformat()
# 根据地址获取tokenzier
def get_tokenizer(tokenizer_path):
    tokenizer = LlamaTokenizer.from_pretrained(tokenizer_path)
    return tokenizer
# 输入文本与tokenizer，返回文本的input_ids
def get_input_ids(text, tokenizer):
    inputs_ids = tokenizer(text).input_ids
    return inputs_ids
# 输入文件地址，拆分地址字符串，返回文件名，文件后缀
def get_filename_suffix(path):
    filename = path.split("/")[-1]
    filename_list = filename.split(".")
    if len(filename_list) == 2:
        suffix = filename_list[-1]
        new_filename = filename_list[0]
    else:
        suffix = filename_list[-1]
        new_filename = ".".join(filename_list[:-1])
    return new_filename, suffix
# 判断文件后缀属于哪种类型
def get_file_type(suffix):
    if suffix == "pdf":
        return "pdf"
    elif suffix == "docx" or suffix == "DOCX":
        return "word"
    elif suffix == "doc" or suffix == "DOC":
        return "doc"
    elif suffix == "txt":
        return "txt"
    elif suffix == "epub":
        return "epub"
    elif suffix == "xlsx":
        return "xlsx"
    elif suffix == "csv":
        return "csv"
    elif suffix == "pptx" or suffix == "ppt":
        return "ppt"
# 判断文件名称是否重复
def is_repeat_filename(filenames, new_filename):
    pattern = r"\(\d{1,3}\)"
    if new_filename in filenames:
        return (1, new_filename)
    
    up_name, down_name = new_filename[:-8], new_filename[-8:]
    match = re.search(pattern, down_name)
    
    if match:
        down_name = down_name.replace(match.group(0), "")
        down_name = down_name.strip()
        final_filename = up_name+down_name
        if final_filename in filenames:
            return (1, final_filename)
        else:
            return (0, final_filename)
    else:
        return (0, new_filename)
# 读取pdf的页数
def get_pdf_pages(path):
    try:
        with pdfplumber.open(path) as pdf:
            page = pdf.pages
            length = len(page)
        print("read {} end!".format(path))
        return length
    except:
        print("read {} error, please check it!".format(path))
        return 0
# 读取单栏显示的pdf文件
def pdf_extract_text_one(path):
    texts = []
    try:
        with pdfplumber.open(path) as pdf:
            page = pdf.pages
            length = len(page)
            for i in tqdm(range(length)):
                text = page[i].extract_text()
                texts.append(text)
        print("read {} end!".format(path))  
        return " ".join(texts), length
    except:
        print("read {} error, please check it!".format(path))
        return path, 0
# 读取双栏显示的pdf文件
def pdf_extract_text_two(pdfPath,language,zoom_x=5,zoom_y=5,rotation_angle=0):
    """_summary_
    Args:
        pdfpath (str): pdf地址
        language (str): 语言
        zoom_x (int, optional): _description_. Defaults to 5.
        zoom_y (int, optional): _description_. Defaults to 5.
        rotation_angle (int, optional): _description_. Defaults to 0.
    Returns:
        _type_: _description_
    """
    try:
        pdf = fitz.open(pdfPath)
        files = os.path.splitext(pdfPath)
        files_list = files[0].split("/")
        filename = files_list[-1]
        pages_length = pdf.page_count
        imgPath = os.path.join("/".join(files_list[:-1]), "images_{}".format(filename))
        if os.path.exists(imgPath):
            shutil.rmtree(imgPath)
            os.mkdir(imgPath)
        else:
            os.mkdir(imgPath)
        # 逐页读取PDF
        for pg in tqdm(range(pages_length)):
            page = pdf[pg]
            # 设置缩放和旋转系数
            trans = fitz.Matrix(zoom_x, zoom_y).prerotate(rotation_angle)
            pm = page.get_pixmap(matrix=trans, alpha=False)
            # 开始写图像
            pm.save(os.path.join(imgPath,str(pg)+".png"))
            #pm.writePNG(imgPath)
        pdf.close()
        
        all_text = []
        for i in tqdm(range(pages_length)):
            img = cv.imread(os.path.join(imgPath,str(i)+".png"))
            if language == "zh":
                boxes = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT, lang='chi_sim+eng')
                text = pytesseract.image_to_string(Image.fromarray(img), lang='chi_sim+eng')
                # text_list = text.split()
                # text = "".join(text_list)
            elif language=="en":
                text = pytesseract.image_to_string(Image.fromarray(img))
            text = text.replace(r"\n\n", r"\n")
            after_text = text.replace(r"\n", "")
        all_text.append(after_text)
        shutil.rmtree(imgPath)
        print("read {} end!".format(pdfPath))
        return " ".join(all_text), pages_length
    except:
        print("read {} error, please check it!".format(pdfPath))
        return pdfPath, 0
# 读取word文件
def docx_extract_text(path):
    try:
        text_list = []
        doc = docx.Document(path)
        length = len(doc.paragraphs)
        for para in tqdm(doc.paragraphs):
            text_list.append(para.text)
        print("read {} end!".format(path))
        return "\n".join(text_list).strip(), length
    except:
        print("read {} error, please check it!".format(path))
        return path, 0 
# 将doc转换成docx
def convert_doc_to_docx(path):
    path_list = path.split("/")
    filename = path_list[-1]
    new_filename = filename.split(".")[0]
    target_dir = "/".join(path_list[:-1])
    os.system(f'soffice --headless --convert-to docx --outdir {target_dir}  {path}')
    return os.path.join(target_dir, new_filename+".docx")
# 将pdf转docx并读取
def convert_new_pdf(path, timeout):
    # 将文件复制到指定目录
    save_dir = "/data/xuhao/xcxhy/Focus_Dataset_v2/pdfs"
    # 从path中解析文件名
    file_path = path.replace("/data/xuhao/xcxhy/Focus_Dataset", "")
    command = f"paddleocr --image_dir={path} --type=structure --recovery=true --output=/data/xuhao/xcxhy/Focus_Dataset_v2/words"
    # 将命令字符串转换为列表形式
    command_list = shlex.split(command)
    return 
# 转换pdf成docx并读取docx内容
def convert_pdf_to_text(path, id):
    # url = 'http://192.168.25.186:7000/convert_pdf'
    url = "http://192.168.25.186:9005/convert_pdf"
    id_str = StringIO(id)
    file = {'file': open(path, 'rb'), "id" : id_str}
    try:
        response = post(url, files=file,timeout=600)
        status_code = response.status_code
        if status_code == 200:
            content = response.json()
            text = content["text"]
        elif status_code == 504:
            text = "read {} timeout!".format(path)
        else:
            text = "can not readed {} !".format(path)
        return text
    except:
        text = "can not read {} !".format(path)
        return text
# 新版本转换pdf成docx并读取docx内容
def new_convert_pdf_to_text(path, id):
    input_path = path.replace("/home/xuhao/xcxhy/", "")
    url = "http://192.168.25.186:9005/convert_pdf"
    data = {"path": input_path, "id" : id}
    try:
        response = post(url, data=json.dumps(data))
        status_code = response.status_code
        if status_code == 200:
            content = response.json()
            text = content["text"]
        elif status_code == 504:
            text = "read {} timeout!".format(path)
        elif status_code == 345:
            text = "json input error {}!".format(path)
        else:
            text = "can not readed {one} and reason by {two}".format(one=path, two=status_code)
        return text
    except:
        text = "ERROR {} !".format(path)
        return text
# 0817最新版本转换pdf成docx并读取docx内容
def new_convert_pdf_to_text_with_split(path, id):
    # 识别pdf得页数
    try:
        pdf_reader = PdfReader(path)
        pages_length = len(pdf_reader.pages)
        if pages_length < 30:
            text = new_convert_pdf_to_text(path, id)
            return text
        else:
            save_dir = "/home/xuhao/xcxhy/Focus_Dataset_v2/pdfs"
            save_path = split_pdf(path, save_dir, id)
            rcp_res = doScp(save_path)
            if rcp_res == 0:
                text = "scp error, please check it!"
                return text 
            else:
                text_list = []
                files = os.listdir(save_path)
                for file in files:
                    file_path = os.path.join(save_path, file)
                    text = new_convert_pdf_to_text(file_path, id)
                    text_list.append(text)
                text = "".join(text_list)
                if "Can not use paddle to" in text:
                    return "this is can not paddle in {} it!".format(path)
                else:
                    return text
    except:
        text = "can not open {} !".format(path)
        return text
# 将66上得文件上传至186
def doScp(path):
    username='root'
    aim_ip='192.168.25.186'
    password='mic@3657'
    source_file_path=path
    aim_file_path='/data/xuhao/xcxhy/Focus_Dataset_v2/pdfs/'
    password_key = '.*assword.*'
    command = f'scp -r {source_file_path}  {username}@{aim_ip}:{aim_file_path}'
    # print("执行指令: ", command)
    try:
        execute = pexpect.spawn(command)
        execute.expect(password_key)
        execute.sendline(password)
        execute.expect(pexpect.EOF)
        return 1
    except:
        return 0
# 切分长度过长得PDF成多个PDF
def split_pdf(pdf_path, save_dir, id):
    # 拆分出来的pdf保存的路径
    if ".pdf" in pdf_path:
        filename, suffix = get_filename_suffix(pdf_path)
        save_path = os.path.join(save_dir, id)
        if not os.path.exists(save_path):
            os.mkdir(save_path)
    pdf_reader = PdfReader(pdf_path)
    pages_length = len(pdf_reader.pages)
    pdfs_nums = int(math.ceil(pages_length/30))
    for i in range(pdfs_nums):
        pdf_writer = PdfWriter()
        # every pdf has 100 pages
        for j in range(i*30, (i+1)*30):
            pdf_writer.add_page(pdf_reader.pages[j])
            if j == pages_length-1:
                break
        with open(os.path.join(save_path,'{}.pdf'.format(str(i))), 'wb') as fh:
            pdf_writer.write(fh)
    return save_path
# 根据sep分割字符串，并设置成固定得长度得文本列表
def slide_window(texts, slide_length, tokenizer, sep):
    text_list = texts.split(sep)
    text_inputs_list = [get_input_ids(text, tokenizer) for text in text_list]
    
    origin_tokens_length = len(sum(text_inputs_list, []))
    slide_tokens_length = 0
    
    slide_text_list = []
    i, j = 0, 1
    while j <= len(text_list):
        choose_input_ids = sum(text_inputs_list[i:j],[])
        text_length = len(choose_input_ids)
        if text_length <= slide_length and j<len(text_list):
            j += 1
        elif text_length <= slide_length and j == len(text_list):
            choose_text = text_list[i:j]
            slide_tokens_length += (len(sum(text_inputs_list[i:j],[])))
            result_text = " ".join(choose_text)
            slide_text_list.append(result_text)
            i = j-2
            j += 1 
        elif j-i <= 2 and text_length > slide_length:
            choose_text = text_list[i:j-1]
            slide_tokens_length += (len(sum(text_inputs_list[i:j-1],[])))
            result_text = " ".join(choose_text)
            slide_text_list.append(result_text)
            i = j-1
            j += 1
        else:
            choose_text = text_list[i:j-1]
            slide_tokens_length += (len(sum(text_inputs_list[i:j-1],[])))
            result_text = " ".join(choose_text)
            slide_text_list.append(result_text)
            i = j-2
    return slide_text_list, origin_tokens_length, slide_tokens_length
# 根据类目名称区分中英文
def get_language(category):
    if "zh" in category:
        language = "zh"
    elif "en" in category:
        language = "en"
    else:
        print("your category is not in zh or en, you should check it!")
        return 
    return language
# 获取唯一的文本id
def get_file_id(n=12):
        """
            Returns:
                ret:random six num and letter
        """
        ret = ""
        for i in range(n):
            num = random.randint(0, 9)
            # num = chr(random.randint(48,57))#ASCII表示数字
            letter = chr(random.randint(97, 122))  # 取小写字母
            Letter = chr(random.randint(65, 90))  # 取大写字母
            s = str(random.choice([num, letter, Letter]))
            ret += s
        return ret
# 根据文件类型，获取文件信息
def get_file_info(type, file_info):
    choose_file = {}
    for id in list(file_info.keys()):
        file_dict = file_info[id]
        if file_dict["file_type"] == type:
            choose_file[file_dict["file_title"]]={"file_pages": file_dict["file_pages"]}
        else:
            pass
    return choose_file
# 根据文件类型，获取文件页数
def get_file_pages(type, file_path):
    if type == "pdf":
        page = get_pdf_pages(file_path)
        return page
    elif type == "word":
        text, page = docx_extract_text(file_path)
        return page
    else:
        return 1

# 过滤已经读取过的ids
def filter_ids(ids, new_ids):
    filter_ids = []
    for id in new_ids:
        if id not in ids:
            filter_ids.append(id)
    return filter_ids
#根据文件id，读取文件信息
def update_file_info(file_ids, file_info):
    update_info = {}
    for id in file_ids:
        update_info[id] = file_info[id]
    return update_info
# 使用正则表达式将中英文单词分开5]+|[a-
def split_string(text):
    pattern = re.compile(r'([\u4e00-\u9fa5]+|[a-zA-Z]+|[^\u4e00-\u9fa5a-zA-Z]+)')
    words = pattern.findall(text)
    return words
# 判断一个字符是否是英文
def is_english(char):
    return 'a' <= char <= 'z' or 'A' <= char <= 'Z' or char == '\s' or char == '\d' or char == '\n' or char
# 判断一个字符串判断是否包含中文字符
def is_contain_chinese(word):
    pattern = re.compile(r'[\u4e00-\u9fa5]')
    match = pattern.search(word)
    return True if match else False
def is_all_only_english(text):
    for char in text:
        if not ('a' <= char <= 'z' or 'A' <= char <= 'Z'):
            return False
    return True

# 判断一个字符串是否全部是英文
def is_all_english(text):
    for char in text:
        if not is_english(char):
            return False
    return True
# 拆分连在一起的英文
def split_wordninja(text):
    text_list = wordninja.split(text)
    return text_list
# 添加空格
def add_space(s):
    new = s
    marks = [",", ";", ":", "!", "?", "，", "。", "？", "！", "；", "：", "、"]
    for i in marks:
        newmark = i + ' '
        ss = new.split(i)
        for j in range(len(ss)):
            ss[j] = ss[j].strip()
        new = newmark.join(ss)
    return new.strip()
# 删除特殊字符得空格
def delete_special_space(text):
    text = text.strip()
    delete_before_list = [",", ";", ":", "!", "?", "，", "。", "？", "！", "；", "：", "、"]
    # delete_after_list = ["(", "[", "{", "<", "（","【", "《", "“", "‘"]
    delete_pairs_list = [("{", "}"), ("[", "]"), ("(", ")"), ("<", ">"), ("\'", "\'"), ("\"", "\""), ("（", "）"), ("【", "】"), ("‘", "’"), ("“", "”"), ("《", "》")]
    # for delete_before in delete_before_list:
    #     text = re.sub("\\s+{}\\s+".format(delete_before), "{} ".format(delete_before), text)
    # for delete_after in delete_after_list:
    #     if delete_after in text:
    #         text = re.sub("{}\s+".format(delete_after), delete_after, text)
    #     else:
    #         pass
    for pairs in delete_pairs_list:
        before, after = pairs
        beforemark, aftermark = ' ' + before, after + ' '
        ss = text.split(before)
        for j in range(len(ss)):
            ss[j] = ss[j].strip()
        text = beforemark.join(ss)
        sd = text.split(after)
        for i in range(len(sd)):
            sd[i] = sd[i].strip()
        text = aftermark.join(sd)
    return text.strip()
# 移除text中得HTML标签
def remove_html_tag(text):
    soup = BeautifulSoup(text, 'html.parser')
    clean_text = soup.get_text()
    return clean_text
# 将二维数组中有交集得list合并
def merge_list_with_intersection(lists):
    # start_list = list(set(list(chain.from_iterable(lists))))
    # print(len(start_list))
    merged = []
    # all_list = []
    for sublist in tqdm(lists):
        sublist = sorted(sublist)
        if len(merged) == 0:
            merged.append(sublist)
            continue
        elif sublist in merged:
            continue
        for i in range(len(merged)):
            merge = set(merged[i])
            if set(sublist).intersection(merge):
                merge_list = list(merge)
                merge_list.extend(sublist)                
                merge = list(set(merge_list))
                merged[i] = sorted(merge)
                break
            elif set(sublist).intersection(merge) != 1 and i == len(merged) - 1:
                merged.append(sublist)
    # end_list = list(set(list(chain.from_iterable(merged))))
    # print(len(end_list))
    return merged
def merge_lists_with_intersection(lists):
    merged = []
    
    while len(lists) > 0:
        current_list = lists.pop(0)
        merged_list = set(current_list)
        
        i = 0
        while i < len(lists):
            if merged_list.intersection(lists[i]):
                merged_list.update(lists.pop(i))
            else:
                i += 1
        
        merged.append(list(merged_list))
    
    return merged
# 根据文件名判断是哪个数据集
def choose_file_category(file_path):
    filename = file_path
    if "trade" in filename:
        return "trade"
        # if "networks" in filename:
        #     return "networks"
        # elif "wiki" in filename:
        #     return "wiki"
        # elif "forum" in filename:
        #     return "forum"
        # elif "webs" in filename:
        #     return "webs"
        # elif "c4" in filename:
        #     return "c4"
        # elif "langchao" in filename:
        #     return "langchao"
        # elif "wudao" in filename:
        #     return "wudao"
        # elif "books" in filename:
        #     return "books"
        # else:
        #     return "customs"
    else:
        return "normal"
    
def store_minhashes(path, minhashes):
    with open(path, "wb") as f:
        pickle.dump(minhashes, f)
    print("END TO STORE {} MINHASHES".format(path))
    return 

def read_minhashes(path):
    with open(path, "rb") as f:
        minhashes = pickle.load(f)
    return minhashes